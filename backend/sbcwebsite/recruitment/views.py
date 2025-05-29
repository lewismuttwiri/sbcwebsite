from django.shortcuts import render, redirect, get_object_or_404
from django.urls import reverse
from django.contrib import messages
from django.http import HttpResponse
from django.core.mail import send_mail
from django.conf import settings
from django.utils import timezone
from django.views.generic import ListView, DetailView
from rest_framework import viewsets, generics, status, permissions
from rest_framework.decorators import api_view, permission_classes, action
from rest_framework.response import Response
from rest_framework.views import APIView
from drf_yasg.utils import swagger_auto_schema
from drf_yasg import openapi
import docx
from docx.shared import Pt, Inches
from io import BytesIO

from sbcapp.models import CustomUser
from .models import JobAdvertisement, JobApplication
from .serializers import JobAdvertisementSerializer, JobApplicationSerializer, JobApplicationUpdateSerializer

def job_success(request):
    return render(request, 'recruitment/job_success.html')

def download_word(request):
    # Create a new Word document
    doc = docx.Document()
    
    # Add a title
    title = doc.add_heading('Job Applications Report', 0)
    
    # Add a subtitle with the current date
    doc.add_paragraph(f'Generated on: {timezone.now().strftime("%Y-%m-%d %H:%M:%S")}')
    
    # Add a line break
    doc.add_paragraph()
    
    # Get all job applications
    applications = JobApplication.objects.all().order_by('-applied_date')
    
    # Group applications by job advertisement
    job_groups = {}
    for app in applications:
        job_title = app.job_advertisement.title
        if job_title not in job_groups:
            job_groups[job_title] = []
        job_groups[job_title].append(app)
    
    # Add each job group to the document
    for job_title, apps in job_groups.items():
        # Add job title as a heading
        doc.add_heading(f'Position: {job_title}', 1)
        doc.add_paragraph(f'Number of Applications: {len(apps)}')
        
        # Add a table for this job's applications
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        
        # Set the header row
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Name'
        header_cells[1].text = 'Email'
        header_cells[2].text = 'Phone'
        header_cells[3].text = 'Status'
        header_cells[4].text = 'Skills'
        header_cells[5].text = 'Application Date'
        
        # Add data rows
        for app in apps:
            row_cells = table.add_row().cells
            row_cells[0].text = app.applicant_name
            row_cells[1].text = app.email
            row_cells[2].text = app.phone
            row_cells[3].text = app.get_status_display()
            row_cells[4].text = ', '.join(app.skills) if app.skills else ''
            row_cells[5].text = app.applied_date.strftime("%Y-%m-%d")
        
        # Add a line break after each table
        doc.add_paragraph()
    
    # Save the document to a BytesIO object
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    # Create the HTTP response with the Word document
    response = HttpResponse(
        buffer.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = 'attachment; filename=job_applications_report.docx'
    
    return response

# API views
class JobAdvertisementViewSet(viewsets.ModelViewSet):
    """
    API endpoint for job advertisements.
    """
    serializer_class = JobAdvertisementSerializer
    
    def get_permissions(self):
        if self.action in ['list', 'retrieve']:
            return [permissions.AllowAny()]
        return [permissions.IsAdminUser()]
    
    def get_queryset(self):
        queryset = JobAdvertisement.objects.all()
        
        # Filter by active status
        active_only = self.request.query_params.get('active_only', 'false').lower() == 'true'
        if active_only:
            queryset = queryset.filter(
                is_active=True,
                closing_date__gte=timezone.now()
            )
        
        # Filter by location
        location = self.request.query_params.get('location')
        if location:
            queryset = queryset.filter(location__icontains=location)
        
        return queryset
    
    @action(detail=True, methods=['get'])
    @swagger_auto_schema(
        operation_description="Get applications for a specific job advertisement (HR or admin only)",
        responses={
            200: openapi.Response(
                description="List of applications for this job",
                schema=JobApplicationSerializer(many=True)
            ),
            403: "Permission denied",
            404: "Not found"
        }
    )
    def applications(self, request, pk=None):
        # Allow HR users or admin users to see applications
        user = request.user
        if not user.is_authenticated or (user.user_role != 4 and not user.is_staff):
            return Response(
                {"error": "You don't have permission to view applications. HR role required."},
                status=status.HTTP_403_FORBIDDEN
            )
        
        job = self.get_object()
        applications = job.applications.all()
        serializer = JobApplicationSerializer(
            applications,
            many=True,
            context={'request': request}
        )
        return Response(serializer.data)

class JobApplicationViewSet(viewsets.ModelViewSet):
    """
    API endpoint for job applications.
    """
    serializer_class = JobApplicationSerializer
    
    def get_permissions(self):
        if self.action == 'create':
            return [permissions.AllowAny()]
        return [permissions.IsAuthenticated()]
    
    def get_queryset(self):
        user = self.request.user
        
        # Allow HR users or admin users to see all applications
        if user.is_authenticated and (user.user_role == 4 or user.is_staff):
            return JobApplication.objects.all()
        
        # For other users, return empty queryset
        return JobApplication.objects.none()
    
    def get_serializer_class(self):
        if self.action in ['update', 'partial_update']:
            return JobApplicationUpdateSerializer
        return JobApplicationSerializer
    
    @swagger_auto_schema(
        operation_description="List all job applications (HR or admin only)",
        responses={
            200: openapi.Response(
                description="List of job applications",
                schema=JobApplicationSerializer(many=True)
            ),
            403: "Permission denied"
        }
    )
    def list(self, request, *args, **kwargs):
        user = request.user
        if not user.is_authenticated or (user.user_role != 4 and not user.is_staff):
            return Response(
                {"error": "You don't have permission to view applications. HR role required."},
                status=status.HTTP_403_FORBIDDEN
            )
        return super().list(request, *args, **kwargs)
    
    @swagger_auto_schema(
        operation_description="Submit a job application",
        request_body=JobApplicationSerializer,
        responses={
            201: openapi.Response(
                description="Job application submitted successfully",
                schema=JobApplicationSerializer()
            ),
            400: "Bad request",
            404: "Job advertisement not found or closed"
        }
    )
    def create(self, request, *args, **kwargs):
        serializer = self.get_serializer(data=request.data)
        serializer.is_valid(raise_exception=True)
        
        # Get the job advertisement
        job_ad_id = request.data.get('job_advertisement')
        try:
            job_ad = JobAdvertisement.objects.get(pk=job_ad_id)
        except JobAdvertisement.DoesNotExist:
            return Response(
                {"error": "Job advertisement not found"},
                status=status.HTTP_404_NOT_FOUND
            )
        
        # Check if the job is still active and accepting applications
        if not job_ad.is_active or job_ad.closing_date < timezone.now():
            return Response(
                {"error": "This job advertisement is no longer accepting applications"},
                status=status.HTTP_400_BAD_REQUEST
            )
        
        # Save the application
        self.perform_create(serializer)
        
        # Send email notification to all HR users
        subject = f"New Job Application: {job_ad.title}"
        message = f"""
        A new job application has been submitted:
        
        Job: {job_ad.title}
        Applicant: {serializer.validated_data.get('applicant_name')}
        Email: {serializer.validated_data.get('email')}
        Phone: {serializer.validated_data.get('phone')}
        
        Please log in to the system to view the full application.
        """
        from_email = settings.DEFAULT_FROM_EMAIL
        
        # Get all HR users' emails
        hr_users = CustomUser.objects.filter(user_role=4, is_active=True)
        recipient_list = [user.email for user in hr_users if user.email]
        
        # Add a fallback email if no HR users are found
        if not recipient_list:
            recipient_list = ['hro@sbckenya.com']
        
        try:
            send_mail(subject, message, from_email, recipient_list)
        except Exception as e:
            # Log the error but don't prevent the application from being submitted
            print(f"Email sending failed: {e}")
        
        headers = self.get_success_headers(serializer.data)
        return Response(
            {
                "message": "Your application has been submitted successfully",
                "application": serializer.data
            },
            status=status.HTTP_201_CREATED,
            headers=headers
        )
    
    @action(detail=False, methods=['get'])
    @swagger_auto_schema(
        operation_description="Get statistics about job applications (HR or admin only)",
        responses={
            200: openapi.Response(
                description="Job application statistics",
                schema=openapi.Schema(
                    type=openapi.TYPE_OBJECT,
                    properties={
                        'total_applications': openapi.Schema(type=openapi.TYPE_INTEGER),
                        'applications_by_job': openapi.Schema(type=openapi.TYPE_OBJECT),
                        'applications_by_status': openapi.Schema(type=openapi.TYPE_OBJECT),
                    }
                )
            ),
            403: "Permission denied"
        }
    )
    def statistics(self, request):
        # Allow HR users or admin users to view statistics
        user = request.user
        if not user.is_authenticated or (user.user_role != 4 and not user.is_staff):
            return Response(
                {"error": "You don't have permission to view statistics. HR role required."},
                status=status.HTTP_403_FORBIDDEN
            )
        
        # Get all applications
        applications = JobApplication.objects.all()
        
        # Total applications
        total_applications = applications.count()
        
        # Applications by job
        from django.db.models import Count
        applications_by_job = JobApplication.objects.values(
            'job_advertisement__title'
        ).annotate(
            count=Count('id')
        ).order_by('-count')
        
        # Applications by status
        applications_by_status = {}
        for status_code, label in JobApplication.STATUS_CHOICES:
            applications_by_status[label] = JobApplication.objects.filter(
                status=status_code
            ).count()
        
        return Response({
            'total_applications': total_applications,
            'applications_by_job': {
                item['job_advertisement__title']: item['count']
                for item in applications_by_job
            },
            'applications_by_status': applications_by_status
        })
