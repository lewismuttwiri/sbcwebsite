from django.shortcuts import render, redirect, get_object_or_404
from django.urls import reverse
from django.contrib import messages
from django.http import HttpResponse
from django.core.mail import send_mail
from django.conf import settings
from django.utils import timezone
from django.views.generic import ListView, DetailView
from .models import JobAdvertisement, JobApplication
from .forms import JobApplicationForm
import docx
from docx.shared import Pt, Inches
from io import BytesIO

class AdvertisementListView(ListView):
    model = JobAdvertisement
    template_name = 'recruitment/advertisement_list.html'
    context_object_name = 'advertisements'
    
    def get_queryset(self):
        return JobAdvertisement.objects.filter(
            is_active=True,
            closing_date__gte=timezone.now()
        )

class AdvertisementDetailView(DetailView):
    model = JobAdvertisement
    template_name = 'recruitment/advertisement_detail.html'
    context_object_name = 'advertisement'

def apply(request, advertisement_id):
    advertisement = get_object_or_404(JobAdvertisement, pk=advertisement_id)
    
    if request.method == 'POST':
        form = JobApplicationForm(request.POST, request.FILES)
        if form.is_valid():
            application = form.save(commit=False)
            application.job_advertisement = advertisement
            application.save()
            
            # Send email notification
            subject = f"New Job Application: {advertisement.title}"
            message = f"""
            A new job application has been submitted:
            
            Job: {advertisement.title}
            Applicant: {application.full_names}
            Email: {application.email}
            Phone: {application.phone_number}
            
            Please log in to the admin panel to view the full application.
            """
            from_email = settings.DEFAULT_FROM_EMAIL
            recipient_list = ['hro@sbckenya.com']  # HR email address
            
            try:
                send_mail(subject, message, from_email, recipient_list)
            except Exception as e:
                # Log the error but don't prevent the application from being submitted
                print(f"Email sending failed: {e}")
            
            return redirect('recruitment:job_success')
    else:
        form = JobApplicationForm()
    
    return render(request, 'recruitment/job_application_form.html', {
        'form': form,
        'advertisement': advertisement
    })

def job_success(request):
    return render(request, 'recruitment/job_success.html')

def job_applied(request):
    applications = JobApplication.objects.all().order_by('-application_date')
    return render(request, 'recruitment/job_applied.html', {'applications': applications})

def download_word(request):
    # Create a new Word document
    doc = docx.Document()
    
    # Add a title
    title = doc.add_heading('Job Applications Report', 0)
    
    # Add a subtitle with the current date
    doc.add_paragraph(f'Generated on: {timezone.now().strftime("%Y-%m-%d %H:%M:%S")}')
    
    # Add a line break
    doc.add_paragraph()
    
    # Get all job applications
    applications = JobApplication.objects.all().order_by('-application_date')
    
    # Group applications by job advertisement
    job_groups = {}
    for app in applications:
        job_title = app.job_advertisement.title
        if job_title not in job_groups:
            job_groups[job_title] = []
        job_groups[job_title].append(app)
    
    # Add each job group to the document
    for job_title, apps in job_groups.items():
        # Add job title as a heading
        doc.add_heading(f'Position: {job_title}', 1)
        doc.add_paragraph(f'Number of Applications: {len(apps)}')
        
        # Add a table for this job's applications
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        
        # Set the header row
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Name'
        header_cells[1].text = 'Email'
        header_cells[2].text = 'Phone'
        header_cells[3].text = 'Education'
        header_cells[4].text = 'Application Date'
        
        # Add data rows
        for app in apps:
            row_cells = table.add_row().cells
            row_cells[0].text = app.full_names
            row_cells[1].text = app.email
            row_cells[2].text = app.phone_number
            row_cells[3].text = f"{app.education_level} - {app.course_of_study}"
            row_cells[4].text = app.application_date.strftime("%Y-%m-%d")
        
        # Add a line break after each table
        doc.add_paragraph()
    
    # Save the document to a BytesIO object
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    # Create the HTTP response with the Word document
    response = HttpResponse(
        buffer.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = 'attachment; filename=job_applications_report.docx'
    
    return response
