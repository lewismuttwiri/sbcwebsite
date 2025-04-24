from django.views.generic import ( TemplateView,)
from django.shortcuts import render, get_object_or_404, redirect
from sbcapp1.models import Product,CartItem,Order,JobAdvertisement,JobApplication,Media,Tender,RelatedImage
from django.views.generic import TemplateView
from django.shortcuts import render
from django.http import JsonResponse
import json
from django.contrib.auth import get_user_model
from django.contrib.auth.decorators import login_required
from django.views.decorators.csrf import csrf_protect
from .forms import JobApplicationForm
from django.http import HttpResponse
from docx import Document
from openpyxl import Workbook
from django.db.models import Sum
from .forms import CommentForm
from django.urls import reverse
from django.utils import timezone
from datetime import timedelta
from django.shortcuts import redirect
from django.contrib.auth import logout


CustomUser = get_user_model()
from allauth.account.views import PasswordResetFromKeyView
from django.http import HttpResponseRedirect

class CustomPasswordResetFromKeyView(PasswordResetFromKeyView):
    def get_success_url(self):
        return reverse('account_login')

def home(request):
    return render(request, 'sbcapp1/home.html')


def home(request):
    media_items = Media.objects.order_by('-datetime_posted')[:3]
    return render(request, 'sbcapp1/home.html', {'media_items': media_items})

class AboutView(TemplateView):
    template_name = "sbcapp1/about.html"

class brandsView(TemplateView):
    template_name = "sbcapp1/brands.html"

def product_list(request):
    products = Product.objects.all()
    return render(request, 'sbcapp1/product_list.html', {'products': products})

def add_to_cart(request, product_id):
    if request.method == 'POST':
        product = get_object_or_404(Product, pk=product_id)
        data = json.loads(request.body)
        quantity = int(data.get('quantity', 1))
        total_price = product.price * quantity
        cart_item, created = CartItem.objects.get_or_create(
            user=request.user,
            product=product,
            defaults={'quantity': quantity, 'total_price': total_price}
        )
        if not created:
            cart_item.quantity += quantity
            cart_item.total_price = product.price * cart_item.quantity
            cart_item.save()
        cart_items_count = CartItem.objects.filter(user=request.user).count()
        return JsonResponse({'message': 'Cart item quantity updated successfully', 'count': cart_items_count})
    else:
        return JsonResponse({'error': 'Method not allowed'}, status=405)

@login_required
def cart_view(request):
    cart_items = CartItem.objects.filter(user=request.user)
    total_price = sum(item.total_price for item in cart_items)
    return render(request, 'sbcapp1/cart.html', {'cart_items': cart_items, 'total_price': total_price})

@login_required
def adjust_quantity(request, item_id):
    cart_item = get_object_or_404(CartItem, id=item_id, user=request.user)
    if request.method == 'POST':
        action = request.POST.get('action')
        if action == 'increment':
            cart_item.quantity += 1
        elif action == 'decrement':
            if cart_item.quantity > 1:
                cart_item.quantity -= 1
        cart_item.total_price = cart_item.product.price * cart_item.quantity
        cart_item.save()
        cart_items = CartItem.objects.filter(user=request.user)
        total_price = sum(item.total_price for item in cart_items)
        return render(request, 'sbcapp1/cart.html', {'cart_items': cart_items, 'total_price': total_price})
    return render(request, 'sbcapp1/cart.html', {'cart_items': [cart_item], 'total_price': cart_item.total_price})

def remove_from_cart(request, item_id):
    cart_item = get_object_or_404(CartItem, id=item_id, user=request.user)
    cart_item.delete()
    return redirect('cart')

def cart_item_count(request):
    if request.user.is_authenticated:
        cart_items_count = CartItem.objects.filter(user=request.user).count()
        return JsonResponse({'count': cart_items_count})
    return JsonResponse({'count': 0})

@login_required
@csrf_protect
def checkout(request):
    if request.method == 'POST':
        name = request.POST.get('name')
        phone_number = request.POST.get('phone_number')
        region = request.POST.get('region')
        cart_items = CartItem.objects.filter(user=request.user)
        total_price = sum(item.total_price for item in cart_items)
        for item in cart_items:
            Order.objects.create(
                user=request.user,
                product_name=item.product.name,
                quantity=item.quantity,
                total_price=item.total_price,
                phone_number=phone_number,
                region=region,
                name=name
            )
        cart_items.delete()
        return redirect('checkout_success')
    else:
        return JsonResponse({'error': 'Method not allowed'}, status=405)

def checkout_success(request):
    return render(request, 'sbcapp1/checkout_success.html')

def order_placed(request):
    if request.user.is_authenticated and request.user.email == 'mt.orders@sbckenya.com':
        orders = Order.objects.all()
        total_price = orders.aggregate(total_price=Sum('total_price'))['total_price']
        return render(request, 'sbcapp1/order_placed.html', {'orders': orders, 'total_price': total_price})
    else:
        return redirect('account_login')

def export_orders_to_excel(request):
    orders = Order.objects.all()
    workbook = Workbook()
    worksheet = workbook.active
    headers = ['User Name', 'User Email', 'Phone Number', 'Region', 'Product Name', 'Quantity', 'Total Price', 'Created At']
    worksheet.append(headers)
    for order in orders:
        created_at = order.created_at.replace(tzinfo=None)
        worksheet.append([
            order.user.get_full_name(),
            order.user.email,
            order.phone_number,
            order.region,
            order.product_name,
            order.quantity,
            order.total_price,
            created_at
        ])
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = 'attachment; filename=orders.xlsx'
    workbook.save(response)
    return response

def advertisement_list(request):
    advertisements = JobAdvertisement.objects.all()
    for advertisement in advertisements:
        advertisement.key_responsibilities = advertisement.key_responsibilities.split('\n')
        advertisement.qualifications = advertisement.qualifications.split('\n')
    return render(request, 'sbcapp1/advertisement_list.html', {'advertisements': advertisements})

@login_required
def apply(request, advertisement_id):
    advertisement = JobAdvertisement.objects.get(id=advertisement_id)
    if request.method == 'POST':
        form = JobApplicationForm(request.POST, request.FILES)
        if form.is_valid():
            application = form.save(commit=False)
            application.advertisement = advertisement
            application.save()
            return redirect('job_success')
    else:
        form = JobApplicationForm()
    return render(request, 'sbcapp1/job_application_form.html', {'form': form, 'advertisement': advertisement})

def job_success(request):
    return render(request, 'sbcapp1/job_success.html')

def job_applied(request):
    if request.user.is_authenticated and request.user.email == 'hro@sbckenya.com':
        job_applications = JobApplication.objects.all()
        return render(request, 'sbcapp1/job_applied.html', {'job_applications': job_applications})
    else:
        return redirect('account_login')

def download_word(request):
    if request.user.is_authenticated and request.user.email == 'hro@sbckenya.com':
        job_applications = JobApplication.objects.all()
        document = Document()
        document.add_heading('Job Applications', level=1)
        for job_application in job_applications:
            document.add_heading(job_application.advertisement.title, level=2)
            document.add_paragraph(f'Full Names: {job_application.full_names}')
            document.add_paragraph(f'Email: {job_application.email}')
            document.add_paragraph(f'Phone Number: {job_application.phone_number}')
            document.add_paragraph(f'Gender: {job_application.get_gender_display()}')  # Use get_<fieldname>_display() for choices
            document.add_paragraph(f'Current Address: {job_application.current_address}')
            document.add_paragraph(f'Date of Birth: {job_application.date_of_birth}')
            document.add_paragraph(f'Education Level: {job_application.get_education_level_display()}')  # Choices field
            document.add_paragraph(f'Institution Name: {job_application.institution_name}')
            document.add_paragraph(f'Course of Study: {job_application.course_of_study}')
            document.add_paragraph(f'Education Start Date: {job_application.education_start_date}')
            document.add_paragraph(f'Education End Date: {job_application.education_end_date}')
            document.add_paragraph(f'Qualification Grade: {job_application.qualification_grade}')
            document.add_paragraph(f'Professional Institution Name: {job_application.professional_institution_name}')
            document.add_paragraph(f'Course Name: {job_application.course_name}')
            document.add_paragraph(f'Professional Start Date: {job_application.professional_start_date}')
            document.add_paragraph(f'Professional End Date: {job_application.professional_end_date}')
            document.add_paragraph(f'Professional Qualification Grade: {job_application.professional_qualification_grade}')
            document.add_paragraph(f'Previous Employer Name: {job_application.previous_employer_name}')
            document.add_paragraph(f'Previous Employer Address: {job_application.previous_employer_address}')
            document.add_paragraph(f'Previous Employer Phone: {job_application.previous_employer_phone}')
            document.add_paragraph(f'Job Title: {job_application.job_title}')
            document.add_paragraph(f'Employment Start Date: {job_application.employment_start_date}')
            document.add_paragraph(f'Employment End Date: {job_application.employment_end_date}')
            document.add_paragraph(f'Employment Duties: {job_application.employment_duties}')
            document.add_paragraph(f'Cover Letter: {job_application.cover_letter}')
            document.add_paragraph(f'Resume: {job_application.resume.url if job_application.resume else "No resume uploaded."}')

        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename=job_applications.docx'
        document.save(response)
        return response
    else:
        return redirect('account_login')

def contact(request):
    if request.method == 'POST':
        form = CommentForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('home')
    else:
        form = CommentForm()
    return render(request, 'sbcapp1/contact.html', {'form': form})




def media_list(request):
    media_items = Media.objects.order_by('-datetime_posted')
    return render(request, 'sbcapp1/media_list.html', {'media_items': media_items})






def media_detail(request, pk):
    media_item = get_object_or_404(Media, pk=pk)
    related_media = Media.objects.exclude(pk=pk)
    related_images = media_item.related_images.all()

    return render(request, 'sbcapp1/media_detail.html', {
        'media_item': media_item,
        'related_media': related_media,
        'related_images': related_images,
    })
def media_summary(request):
    media_items = Media.objects.order_by('-datetime_posted')[:2]
    data = [{
        'image': media_item.image.url,
        'title': media_item.title,
        'datetime_posted': media_item.datetime_posted.strftime('%Y-%m-%d %H:%M:%S'),
    } for media_item in media_items]
    return JsonResponse(data, safe=False)

def tender_list(request):
    tenders = Tender.objects.all()
    return render(request, 'sbcapp1/tenders.html', {'tenders': tenders})

def social(request):
    return render(request, 'sbcapp1/social.html')

def privacy_and_terms(request):
    return render(request, 'sbcapp1/privacy.html')

